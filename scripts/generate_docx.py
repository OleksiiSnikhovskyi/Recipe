#!/usr/bin/env python3
"""
generate_docx.py

Generate formatted Microsoft Word (.docx) documents from recipe JSON or a
recipe stored in PostgreSQL.

Usage:
    python scripts/generate_docx.py --recipe recipe.json --output output.docx
    python scripts/generate_docx.py --recipe-id 12 --output-dir output/docx

HTTP:
    POST /generate
    {"recipe_id": 12}
    {"video_id": "abc123"}
    {"recipe": {...recipe JSON...}}
"""

import argparse
import base64
import json
import os
import re
import sys
import requests
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, Optional

from dotenv import load_dotenv

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("Error: python-docx not installed. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


load_dotenv()

DOCX_OUTPUT_DIR = Path(os.getenv("DOCX_OUTPUT_DIR", "output/docx"))
DOCX_PUBLIC_BASE_URL = os.getenv("DOCX_PUBLIC_BASE_URL", "")

DB_HOST = os.getenv("DB_HOST", "127.0.0.1")
DB_PORT = int(os.getenv("DB_PORT", "5432"))
DB_NAME = os.getenv("DB_NAME", "recipe_db")
DB_USER = os.getenv("DB_USER", "recipe_user")
DB_PASSWORD = os.getenv("DB_PASSWORD") or os.getenv("RECIPE_DB_PASSWORD", "")


def _coerce_json(value: Any, fallback: Any) -> Any:
    if value is None:
        return fallback
    if isinstance(value, (dict, list)):
        return value
    if isinstance(value, str):
        try:
            return json.loads(value)
        except json.JSONDecodeError:
            return fallback
    return fallback


def _as_list(value: Any) -> list:
    value = _coerce_json(value, value)
    return value if isinstance(value, list) else []


def _as_dict(value: Any) -> dict:
    value = _coerce_json(value, value)
    return value if isinstance(value, dict) else {}


def safe_filename(value: str, suffix: str = ".docx") -> str:
    value = (value or "recipe").strip()
    value = re.sub(r'[<>:"/\\|?*\x00-\x1F]', " ", value)
    value = re.sub(r"\s+", " ", value).strip(" .")
    if not value:
        value = "recipe"
    return value[:140] + suffix


def format_minutes(value: Any) -> str:
    if value in (None, "", 0):
        return "не вказано"
    try:
        value = int(value)
    except (TypeError, ValueError):
        return str(value)
    hours, minutes = divmod(value, 60)
    if hours and minutes:
        return f"{hours} год {minutes} хв"
    if hours:
        return f"{hours} год"
    return f"{minutes} хв"


def recipe_from_db_row(row: Dict[str, Any]) -> Dict[str, Any]:
    recipe_text = _as_dict(row.get("recipe_text"))
    recipe = dict(recipe_text)
    recipe.update({
        "id": row.get("id"),
        "title": row.get("title") or recipe.get("title") or "Рецепт",
        "category": row.get("category") or recipe.get("category") or "Інше",
        "description": row.get("description") or recipe.get("description") or "",
        "ingredients": _as_list(row.get("ingredients") or recipe.get("ingredients")),
        "steps": _as_list(row.get("steps") or recipe.get("steps")),
        "nutrition": _as_dict(row.get("nutrition") or recipe.get("nutrition")),
        "transcription": recipe.get("transcription") or {
            "source": row.get("transcript_source") or "",
            "language": row.get("transcript_language") or "",
            "warning": row.get("transcription_warning") or "",
        },
    })
    recipe["source"] = {
        **_as_dict(recipe.get("source")),
        "video_id": row.get("video_id") or _as_dict(recipe.get("source")).get("video_id"),
        "video_url": row.get("youtube_url") or _as_dict(recipe.get("source")).get("video_url"),
        "youtube_channel": row.get("youtube_channel") or _as_dict(recipe.get("source")).get("youtube_channel"),
        "thumbnail_url": row.get("thumbnail_url") or _as_dict(recipe.get("source")).get("thumbnail_url"),
    }
    recipe["metadata"] = {
        **_as_dict(recipe.get("metadata")),
        "recipe_id": row.get("id"),
        "video_id": row.get("video_id"),
    }
    return recipe


def db_connection():
    import psycopg2
    import psycopg2.extras

    return psycopg2.connect(
        host=DB_HOST,
        port=DB_PORT,
        dbname=DB_NAME,
        user=DB_USER,
        password=DB_PASSWORD,
        cursor_factory=psycopg2.extras.RealDictCursor,
    )


def fetch_recipe(recipe_id: Optional[int] = None, video_id: Optional[str] = None) -> Dict[str, Any]:
    if not recipe_id and not video_id:
        raise ValueError("recipe_id or video_id is required")

    where = "id = %s" if recipe_id else "video_id = %s"
    value = recipe_id if recipe_id else video_id
    with db_connection() as conn:
        with conn.cursor() as cursor:
            cursor.execute(f"SELECT * FROM recipes WHERE {where} LIMIT 1", (value,))
            row = cursor.fetchone()
    if not row:
        raise LookupError(f"Recipe not found: {value}")
    return recipe_from_db_row(dict(row))


def update_docx_path(recipe_id: Optional[int], video_id: Optional[str], docx_path: str) -> None:
    if not recipe_id and not video_id:
        return
    where = "id = %s" if recipe_id else "video_id = %s"
    value = recipe_id if recipe_id else video_id
    with db_connection() as conn:
        with conn.cursor() as cursor:
            cursor.execute(
                f"UPDATE recipes SET docx_path = %s, updated_at = NOW() WHERE {where}",
                (docx_path, value),
            )


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(86, 64, 42)
    return paragraph


def add_kv(paragraph, label: str, value: Any) -> None:
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(str(value or "не вказано"))


def iter_ingredients(ingredients: Iterable[Dict[str, Any]]) -> Iterable[tuple]:
    for item in ingredients:
        if not isinstance(item, dict):
            yield str(item), "", "", ""
            continue
        yield (
            item.get("name", ""),
            item.get("quantity", ""),
            item.get("unit", ""),
            item.get("notes", ""),
        )


def _nutrition_has_values(nutrition: Dict[str, Any]) -> bool:
    for key in ("per_100g", "per_serving"):
        values = _as_dict(nutrition.get(key))
        if any(values.get(name) not in (None, "") for name in ("calories", "protein", "fat", "carbohydrates")):
            return True
    return False


def _add_thumbnail(doc: Document, thumbnail_url: str) -> None:
    if not thumbnail_url:
        return
    try:
        response = requests.get(thumbnail_url, timeout=20)
        response.raise_for_status()
        image = BytesIO(response.content)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image, width=Inches(5.4))
    except Exception as exc:
        warning = doc.add_paragraph()
        warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
        warning_run = warning.add_run(f"Фото з YouTube не вдалося завантажити: {exc}")
        warning_run.italic = True
        warning_run.font.size = Pt(9)


def _add_video_qr(doc: Document, video_url: str) -> bool:
    if not video_url:
        return False
    try:
        import qrcode

        image = qrcode.make(video_url)
        buffer = BytesIO()
        image.save(buffer, format="PNG")
        buffer.seek(0)

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = paragraph.add_run("QR-код відео")
        caption.bold = True

        qr_paragraph = doc.add_paragraph()
        qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qr_paragraph.add_run().add_picture(buffer, width=Inches(1.35))
        return True
    except Exception:
        return False


def generate_docx(recipe: Dict[str, Any]) -> bytes:
    """Generate DOCX document from recipe data."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = recipe.get("title") or "Рецепт"
    category = recipe.get("category") or "Інше"
    source = _as_dict(recipe.get("source"))
    nutrition = _as_dict(recipe.get("nutrition"))
    ingredients = _as_list(recipe.get("ingredients"))
    steps = _as_list(recipe.get("steps"))

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(86, 64, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(category)
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(128, 96, 58)

    summary = doc.add_paragraph()
    pieces = [
        f"Порцій: {recipe.get('servings') or 'не вказано'}",
        f"Підготовка: {format_minutes(recipe.get('prep_time_minutes'))}",
        f"Приготування: {format_minutes(recipe.get('cook_time_minutes'))}",
    ]
    summary.add_run(" | ".join(pieces))
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_thumbnail(doc, source.get("thumbnail_url", ""))

    if recipe.get("description"):
        add_heading(doc, "Опис", 2)
        doc.add_paragraph(str(recipe["description"]))

    add_heading(doc, "Інгредієнти", 2)
    if ingredients:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["Інгредієнт", "Кількість", "Одиниця", "Примітки"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for name, quantity, unit, notes in iter_ingredients(ingredients):
            row = table.add_row().cells
            row[0].text = str(name or "")
            row[1].text = "" if quantity is None else str(quantity)
            row[2].text = str(unit or "")
            row[3].text = str(notes or "")
    else:
        doc.add_paragraph("Інгредієнти не вказані.")

    add_heading(doc, "Приготування", 2)
    if steps:
        for index, step in enumerate(steps, start=1):
            if isinstance(step, dict):
                number = step.get("step_number") or index
                instruction = step.get("instruction") or ""
                duration = step.get("duration_minutes")
                notes = step.get("notes")
            else:
                number, instruction, duration, notes = index, str(step), None, None
            paragraph = doc.add_paragraph(style=None)
            number_run = paragraph.add_run(f"{number}. ")
            number_run.bold = True
            paragraph.add_run(str(instruction))
            if duration:
                paragraph.add_run(f" ({format_minutes(duration)})").italic = True
            if notes:
                note = doc.add_paragraph(f"Порада: {notes}")
                note.paragraph_format.left_indent = Inches(0.25)
    else:
        doc.add_paragraph("Кроки приготування не вказані.")

    if nutrition and _nutrition_has_values(nutrition):
        add_heading(doc, "Поживність", 2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for index, header in enumerate(["Показник", "Ккал", "Білки", "Жири", "Вуглеводи"]):
            table.rows[0].cells[index].text = header
        labels = [("per_100g", "На 100 г"), ("per_serving", "На порцію")]
        for key, label in labels:
            values = _as_dict(nutrition.get(key))
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(values.get("calories", ""))
            row[2].text = str(values.get("protein", ""))
            row[3].text = str(values.get("fat", ""))
            row[4].text = str(values.get("carbohydrates", ""))

    add_heading(doc, "Джерело", 2)
    source_paragraph = doc.add_paragraph()
    add_kv(source_paragraph, "YouTube канал", source.get("youtube_channel") or "Unknown")
    if source.get("video_url"):
        if not _add_video_qr(doc, source["video_url"]):
            link_paragraph = doc.add_paragraph()
            add_kv(link_paragraph, "Відео", source["video_url"])
    transcription = _as_dict(recipe.get("transcription"))
    if transcription.get("warning"):
        warning = doc.add_paragraph()
        warning.add_run("Попередження транскрипції: ").bold = True
        warning.add_run(str(transcription["warning"]))

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    footer_run = footer.add_run(f"Згенеровано: {generated_at}")
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def write_docx(recipe: Dict[str, Any], output_dir: Path = DOCX_OUTPUT_DIR) -> Dict[str, Any]:
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = safe_filename(recipe.get("title", "recipe"))
    output_path = output_dir / filename
    docx_bytes = generate_docx(recipe)
    output_path.write_bytes(docx_bytes)

    public_url = ""
    if DOCX_PUBLIC_BASE_URL:
        public_url = DOCX_PUBLIC_BASE_URL.rstrip("/") + "/" + filename

    return {
        "docx_bytes": docx_bytes,
        "filename": filename,
        "docx_path": str(output_path),
        "docx_url": public_url,
    }


def create_flask_app():
    try:
        from flask import Flask, jsonify, request
    except ImportError:
        return None

    app = Flask(__name__)

    @app.route("/generate", methods=["POST"])
    def generate_endpoint():
        try:
            data = request.get_json(silent=True) or {}
            recipe = data.get("recipe")
            recipe_id = data.get("recipe_id")
            video_id = data.get("video_id")

            if recipe is None:
                recipe = fetch_recipe(recipe_id=recipe_id, video_id=video_id)
            if not isinstance(recipe, dict):
                return jsonify({"error": "recipe must be an object"}), 400

            generated = write_docx(recipe)
            resolved_recipe_id = recipe.get("id") or _as_dict(recipe.get("metadata")).get("recipe_id") or recipe_id
            resolved_video_id = _as_dict(recipe.get("source")).get("video_id") or _as_dict(recipe.get("metadata")).get("video_id") or video_id
            update_docx_path(resolved_recipe_id, resolved_video_id, generated["docx_path"])

            return jsonify({
                "ok": True,
                "recipe_id": resolved_recipe_id,
                "video_id": resolved_video_id,
                "title": recipe.get("title"),
                "category": recipe.get("category"),
                "filename": generated["filename"],
                "docx_path": generated["docx_path"],
                "docx_url": generated["docx_url"],
                "docx_base64": base64.b64encode(generated["docx_bytes"]).decode("ascii"),
            }), 200
        except Exception as exc:
            return jsonify({"ok": False, "error": str(exc)}), 500

    @app.route("/health", methods=["GET"])
    def health():
        return jsonify({
            "status": "ok",
            "output_dir": str(DOCX_OUTPUT_DIR),
            "db_host": DB_HOST,
            "db_name": DB_NAME,
        }), 200

    return app


def load_recipe_file(path: str) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8") as file:
        return json.load(file)


def main():
    parser = argparse.ArgumentParser(description="Generate DOCX from recipe JSON")
    parser.add_argument("--recipe", "-r", help="Recipe JSON file path")
    parser.add_argument("--recipe-id", type=int, help="Recipe database ID")
    parser.add_argument("--video-id", help="YouTube video ID")
    parser.add_argument("--output", "-o", help="Output DOCX file path")
    parser.add_argument("--output-dir", default=str(DOCX_OUTPUT_DIR), help="Output directory for generated DOCX")
    parser.add_argument("--server", action="store_true", help="Run as Flask server")
    parser.add_argument("--port", default=5001, type=int, help="Server port (default: 5001)")

    args = parser.parse_args()

    if args.server:
        app = create_flask_app()
        if not app:
            print("Error: Flask not installed. Install with: pip install flask", file=sys.stderr)
            sys.exit(1)
        print(f"Starting DOCX Flask server on port {args.port}...")
        app.run(host="0.0.0.0", port=args.port, debug=False)
        return

    if args.recipe:
        recipe = load_recipe_file(args.recipe)
    else:
        recipe = fetch_recipe(recipe_id=args.recipe_id, video_id=args.video_id)

    if args.output:
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(generate_docx(recipe))
        print(f"Generated: {output_path}")
    else:
        result = write_docx(recipe, Path(args.output_dir))
        print(json.dumps({
            "filename": result["filename"],
            "docx_path": result["docx_path"],
            "docx_url": result["docx_url"],
        }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
